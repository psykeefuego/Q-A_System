import os
import json
import time
import datetime
import streamlit as st
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangChainDocument
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import google.generativeai as genai

# Load .env
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# --- Portable FAISS folder ---
FAISS_FOLDER = os.path.join(os.getcwd(), "faiss_index")

# --- Extractors ---
def extract_pdf(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text

def extract_docx(docx_file):
    doc = DocxDocument(docx_file)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def extract_txt(txt_file):
    text = txt_file.read().decode('utf-8')
    return text

def extract_html(html_file):
    soup = BeautifulSoup(html_file, 'html.parser')
    return soup.get_text()

def extract_md(md_file):
    html = markdown.markdown(md_file.read().decode('utf-8'))
    soup = BeautifulSoup(html, 'html.parser')
    return soup.get_text()

def get_file_text(files):
    text = ""
    for file in files:
        filename = file.name.lower()
        if filename.endswith(".pdf"):
            text += extract_pdf(file)
        elif filename.endswith(".docx"):
            text += extract_docx(file)
        elif filename.endswith(".txt"):
            text += extract_txt(file)
        elif filename.endswith(".html"):
            text += extract_html(file)
        elif filename.endswith(".md"):
            text += extract_md(file)
        else:
            st.warning(f"Unsupported file type: {filename}")
    return text

# --- Chunk ---
def get_text_chunks(text, filename):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(text)
    now = datetime.datetime.now().isoformat()
    docs = []
    for chunk in chunks:
        docs.append(LangChainDocument(
            page_content=chunk,
            metadata={"source": filename, "timestamp": now}
        ))
    return docs

# --- Save Vectors ---
def save_vector_store(docs):
    # Embedding model
    embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")

    texts = [doc.page_content for doc in docs]
    metadatas = [doc.metadata for doc in docs]
    vector_store = FAISS.from_texts(texts, embeddings, metadatas=metadatas)
    vector_store.save_local(FAISS_FOLDER)
    return vector_store

# --- RAG Chain ---
def get_conversational_chain(chat_history):
    context_messages = "\n".join(
        [f"{m['role']}: {m['content']}" for m in chat_history[-5:]]
    )
    prompt_template = f"""
    Here is the conversation so far:
    {context_messages}

    Answer the question succinctly using only the retrieved context. Provide only the fact, no extra explanation."

    Context: {{context}}

    Question: {{question}}

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-pro")
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

# --- Feedback save ---
def save_feedback():
    with open("feedback.json", "w") as f:
        json.dump(st.session_state.feedback, f, indent=2)

# --- Eval ---
def simple_eval(pred, true):
    pred_tokens = pred.lower().split()
    true_tokens = true.lower().split()
    common = set(pred_tokens) & set(true_tokens)
    precision = len(common) / len(pred_tokens) if pred_tokens else 0
    recall = len(common) / len(true_tokens) if true_tokens else 0
    if precision + recall == 0:
        return 0, 0, 0
    f1 = 2 * (precision * recall) / (precision + recall)
    return precision, recall, f1

# --- Load SQuAD sample ---
def load_squad_sample():
    with open("squad_sample.json") as f:
        samples = json.load(f)
    docs = []
    now = datetime.datetime.now().isoformat()
    for s in samples:
        docs.append(LangChainDocument(
            page_content=s["context"],
            metadata={"source": "squad_sample", "timestamp": now}
        ))
    save_vector_store(docs)
    return samples

# --- User input ---
def user_input(user_question):
    if st.session_state.request_count >= 10:
        st.warning("Rate limit reached.")
        return

    st.session_state.request_count += 1

    if user_question in st.session_state.answer_cache:
        st.write("Cached Reply:", st.session_state.answer_cache[user_question])
        return

    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local(
        FAISS_FOLDER,
        embeddings,
        allow_dangerous_deserialization=True 
    )

    start_time = time.time()

    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain(st.session_state.chat_history)

    try:
        response = chain(
            {"input_documents": docs, "question": user_question},
            return_only_outputs=True
        )
    except Exception as e:
        st.error(f"Error: {str(e)}")
        with open("logs.txt", "a") as f:
            f.write(f"{datetime.datetime.now()} ERROR: {str(e)}\n")
        return

    duration = time.time() - start_time

    st.session_state.chat_history.append({"role": "assistant", "content": response["output_text"]})
    st.session_state.answer_cache[user_question] = response["output_text"]

    st.write("Reply:", response["output_text"])
    st.write(f"Response time: {duration:.2f} sec")

    st.subheader("Feedback-")
    if st.button("👍"):
        st.session_state.feedback.append({
            "question": user_question,
            "answer": response["output_text"],
            "feedback": "upvote"
        })
        save_feedback()
        st.success("Thanks for your feedback!")
    if st.button("👎"):
        st.session_state.feedback.append({
            "question": user_question,
            "answer": response["output_text"],
            "feedback": "downvote"
        })
        correction = st.text_area("Your correction:")
        if st.button("Submit Correction"):
            st.session_state.feedback.append({
                "question": user_question,
                "answer": response["output_text"],
                "correction": correction
            })
            save_feedback()
            st.success("Correction saved!")
    rating = st.slider("Rate (1–5):", 1, 5)
    if st.button("Submit Rating"):
        st.session_state.feedback.append({
            "question": user_question,
            "answer": response["output_text"],
            "rating": rating
        })
        save_feedback()
        st.success(f"Thanks for rating: {rating}")

# --- Main ---
def main():
    st.set_page_config("QA System")
    st.header("Intelligent Document Q&A System with Memory")

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "feedback" not in st.session_state:
        st.session_state.feedback = []
    if "answer_cache" not in st.session_state:
        st.session_state.answer_cache = {}
    if "request_count" not in st.session_state:
        st.session_state.request_count = 0

    user_question = st.text_input("Ask a question:")

    if user_question:
        st.session_state.chat_history.append({"role": "user", "content": user_question})
        user_input(user_question)

    with st.sidebar:
        st.title("📂 Upload & Test")
        files = st.file_uploader("Upload docs", accept_multiple_files=True)
        if st.button("Submit & Process"):
            if files:
                raw_text = get_file_text(files)
                filename = files[0].name if files else "unknown"
                docs = get_text_chunks(raw_text, filename)
                save_vector_store(docs)
                st.success(f" Processed {len(files)} files.")
            else:
                st.warning("Upload a file first!")

        if st.button("Load SQuAD Sample"):
            squad_samples = load_squad_sample()
            st.session_state.squad_samples = squad_samples
            st.success("SQuAD sample loaded!")

        if st.button("Run SQuAD Test"):
            if "squad_samples" not in st.session_state:
                st.warning("Load SQuAD first.")
            else:
                embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
                new_db = FAISS.load_local(
                    FAISS_FOLDER,
                    embeddings,
                    allow_dangerous_deserialization=True
                )
                chain = get_conversational_chain([])

                for s in st.session_state.squad_samples:
                    docs = new_db.similarity_search(s["question"])
                    response = chain(
                        {"input_documents": docs, "question": s["question"]},
                        return_only_outputs=True
                    )
                    pred = response["output_text"]
                    p, r, f1 = simple_eval(pred, s["answer"])
                    st.write(f"Q: {s['question']}")
                    st.write(f"Expected: {s['answer']}")
                    st.write(f"Got: {pred}")
                    st.write(f"P: {p:.2f} R: {r:.2f} F1: {f1:.2f}")

if __name__ == "__main__":
    main()
